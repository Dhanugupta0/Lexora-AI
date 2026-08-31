"""Document -> clean, structured text blocks.

The old parser flattened everything into a wall of text, which threw away the
two signals retrieval cares about most: **where** a passage sits in the document
and **what section** it belongs to. This one keeps both.

For each supported format it produces `Block`s tagged with a page number, a
kind (heading / body / table / list) and the heading path they live under.
Along the way it strips repeated running headers and footers, repairs
line-break hyphenation, and normalises whitespace -- all noise that otherwise
ends up polluting embeddings.
"""

from __future__ import annotations

import io
import logging
import re
from collections import Counter
from dataclasses import dataclass, field
from typing import List, Optional, Tuple

logger = logging.getLogger(__name__)

SUPPORTED_TYPES = {"pdf", "docx", "txt", "md", "markdown"}


@dataclass
class Block:
    text: str
    page: int
    kind: str = "body"                    # heading | body | table | list
    level: int = 0                        # heading depth (0 = not a heading)
    section: str = ""                     # breadcrumb of enclosing headings


@dataclass
class ParsedDocument:
    blocks: List[Block] = field(default_factory=list)
    title: str = ""
    page_count: int = 0
    meta: dict = field(default_factory=dict)

    @property
    def pages(self) -> List[str]:
        """Flat per-page text -- kept for backwards compatibility."""
        buckets: dict = {}
        for block in self.blocks:
            buckets.setdefault(block.page, []).append(block.text)
        return ["\n\n".join(v) for _, v in sorted(buckets.items())]

    @property
    def char_count(self) -> int:
        return sum(len(b.text) for b in self.blocks)


# --------------------------------------------------------------------------- #
# Text hygiene
# --------------------------------------------------------------------------- #
_HYPHEN_BREAK = re.compile(r"(\w)-\n(\w)")
_MULTI_NEWLINE = re.compile(r"\n{3,}")
_MULTI_SPACE = re.compile(r"[ \t]{2,}")
_CONTROL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")
_BULLET = re.compile(r"^\s*([-*•●·]|\d+[.)])\s+")


def clean_text(text: str) -> str:
    text = _CONTROL.sub("", text)
    text = _HYPHEN_BREAK.sub(r"\1\2", text)          # re-join words split across lines
    text = text.replace("­", "").replace("ﬁ", "fi").replace("ﬂ", "fl")
    text = _MULTI_SPACE.sub(" ", text)
    text = _MULTI_NEWLINE.sub("\n\n", text)
    return "\n".join(line.rstrip() for line in text.splitlines()).strip()


def _strip_running_headers(pages: List[str], threshold: float = 0.6) -> List[str]:
    """Drop short first/last lines that repeat across most pages."""
    if len(pages) < 4:
        return pages

    edges: Counter = Counter()
    for page in pages:
        lines = [line.strip() for line in page.splitlines() if line.strip()]
        for line in lines[:1] + lines[-1:]:
            if len(line) < 90:
                edges[re.sub(r"\d+", "#", line)] += 1

    cutoff = max(3, int(len(pages) * threshold))
    repeated = {line for line, count in edges.items() if count >= cutoff}
    if not repeated:
        return pages

    cleaned = []
    for page in pages:
        lines = page.splitlines()
        while lines and re.sub(r"\d+", "#", lines[0].strip()) in repeated:
            lines.pop(0)
        while lines and re.sub(r"\d+", "#", lines[-1].strip()) in repeated:
            lines.pop()
        cleaned.append("\n".join(lines))
    return cleaned


def _assign_sections(blocks: List[Block]) -> List[Block]:
    """Walk the block list and stamp each one with its heading breadcrumb."""
    stack: List[Tuple[int, str]] = []
    for block in blocks:
        if block.kind == "heading":
            while stack and stack[-1][0] >= block.level:
                stack.pop()
            block.section = " > ".join(title for _, title in stack)
            stack.append((block.level, block.text.strip()))
        else:
            block.section = " > ".join(title for _, title in stack)
    return blocks


# --------------------------------------------------------------------------- #
# PDF
# --------------------------------------------------------------------------- #
def _parse_pdf(data: bytes) -> ParsedDocument:
    try:
        import pymupdf as fitz          # PyMuPDF >= 1.24.3
    except ImportError:
        import fitz                     # older releases

    raw_pages: List[str] = []
    spans: List[Tuple[int, str, float, bool]] = []      # page, text, size, bold
    tables: List[Block] = []
    title = ""

    with fitz.open(stream=data, filetype="pdf") as doc:
        title = (doc.metadata or {}).get("title") or ""
        page_count = doc.page_count

        for index, page in enumerate(doc, start=1):
            raw_pages.append(page.get_text("text"))

            layout = page.get_text("dict")
            for area in layout.get("blocks", []):
                if area.get("type") != 0:
                    continue
                for line in area.get("lines", []):
                    text = "".join(s.get("text", "") for s in line.get("spans", [])).strip()
                    if not text:
                        continue
                    first = line["spans"][0]
                    bold = bool(first.get("flags", 0) & (1 << 4)) or "bold" in first.get("font", "").lower()
                    spans.append((index, text, round(first.get("size", 0), 1), bold))

            try:
                for table in page.find_tables():
                    md = _table_to_markdown(table.extract())
                    if md:
                        tables.append(Block(text=md, page=index, kind="table"))
            except Exception:                              # noqa: BLE001 -- tables are best-effort
                pass

    if not any(p.strip() for p in raw_pages):
        raise ValueError("PDF has no extractable text (it may be a scanned image).")

    raw_pages = _strip_running_headers(raw_pages)
    body_size = _dominant_size(spans)
    dropped = {re.sub(r"\d+", "#", line) for line in _repeated_lines(raw_pages)}

    blocks: List[Block] = []
    for page_index, page_text in enumerate(raw_pages, start=1):
        page_spans = {text for p, text, _, _ in spans if p == page_index}
        buffer: List[str] = []

        for line in clean_text(page_text).splitlines():
            stripped = line.strip()
            if not stripped or re.sub(r"\d+", "#", stripped) in dropped:
                continue

            level = _heading_level(stripped, page_spans, spans, page_index, body_size)
            if level:
                if buffer:
                    blocks.append(Block(text="\n".join(buffer), page=page_index))
                    buffer = []
                blocks.append(Block(text=stripped, page=page_index, kind="heading", level=level))
            else:
                buffer.append(stripped)

        if buffer:
            blocks.append(Block(text="\n".join(buffer), page=page_index))

    blocks.extend(tables)
    blocks.sort(key=lambda b: (b.page, b.kind == "table"))
    blocks = _assign_sections([b for b in blocks if b.text.strip()])

    if not title:
        heading = next((b.text for b in blocks if b.kind == "heading"), "")
        title = heading or (blocks[0].text[:80] if blocks else "")

    return ParsedDocument(blocks=blocks, title=title.strip(), page_count=page_count,
                          meta={"format": "pdf", "tables": len(tables)})


def _dominant_size(spans: List[Tuple[int, str, float, bool]]) -> float:
    if not spans:
        return 0.0
    weighted: Counter = Counter()
    for _, text, size, _ in spans:
        weighted[size] += len(text)
    return weighted.most_common(1)[0][0]


def _repeated_lines(pages: List[str], threshold: float = 0.6) -> List[str]:
    counts: Counter = Counter()
    for page in pages:
        for line in {l.strip() for l in page.splitlines() if 0 < len(l.strip()) < 90}:
            counts[re.sub(r"\d+", "#", line)] += 1
    cutoff = max(3, int(len(pages) * threshold))
    return [line for line, count in counts.items() if count >= cutoff and len(pages) >= 4]


def _heading_level(line: str, page_spans, spans, page: int, body_size: float) -> int:
    """Heading if the font is visibly larger than body text, or it looks like one."""
    if len(line) > 120 or line.endswith((".", ",", ";")):
        return 0

    if line in page_spans and body_size:
        size = max((s for p, t, s, _ in spans if p == page and t == line), default=0.0)
        if size >= body_size * 1.45:
            return 1
        if size >= body_size * 1.18:
            return 2

    if re.match(r"^\d+(\.\d+)*\.?\s+\S", line):                     # "3.1 Methodology"
        return min(3, line.split()[0].count(".") + 1)
    if re.match(r"^(chapter|section|part|appendix)\b", line, re.I):
        return 1
    if len(line) < 70 and line.isupper() and len(line.split()) <= 10:
        return 2
    return 0


def _table_to_markdown(rows) -> str:
    rows = [[(cell or "").replace("\n", " ").strip() for cell in row] for row in rows if row]
    rows = [row for row in rows if any(row)]
    if len(rows) < 2:
        return ""
    header, *body = rows
    out = ["| " + " | ".join(header) + " |", "|" + "|".join(["---"] * len(header)) + "|"]
    out += ["| " + " | ".join(row + [""] * (len(header) - len(row))) + " |" for row in body[:60]]
    return "\n".join(out)


# --------------------------------------------------------------------------- #
# DOCX
# --------------------------------------------------------------------------- #
def _parse_docx(data: bytes) -> ParsedDocument:
    from docx import Document

    doc = Document(io.BytesIO(data))
    blocks: List[Block] = []
    buffer: List[str] = []
    page = 1
    chars = 0

    def flush():
        nonlocal buffer
        if buffer:
            blocks.append(Block(text="\n".join(buffer), page=page))
            buffer = []

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            flush()
            continue

        style = (para.style.name or "").lower()
        if style.startswith("heading") or style == "title":
            flush()
            digits = re.findall(r"\d+", style)
            level = 1 if style == "title" else int(digits[0]) if digits else 1
            blocks.append(Block(text=text, page=page, kind="heading", level=min(level, 4)))
        else:
            kind = "list" if style.startswith("list") or _BULLET.match(text) else "body"
            if kind == "list" and buffer:
                buffer.append(text)
            else:
                buffer.append(text)

        chars += len(text)
        if chars > 2600:                     # rough page proxy -- docx has no page breaks
            flush()
            page += 1
            chars = 0
    flush()

    for table in doc.tables:
        md = _table_to_markdown([[cell.text for cell in row.cells] for row in table.rows])
        if md:
            blocks.append(Block(text=md, page=page, kind="table"))

    if not blocks:
        raise ValueError("DOCX has no extractable text.")

    blocks = _assign_sections(blocks)
    props = doc.core_properties
    title = (props.title or "").strip() or next((b.text for b in blocks if b.kind == "heading"), "")
    return ParsedDocument(blocks=blocks, title=title, page_count=page,
                          meta={"format": "docx", "tables": len(doc.tables)})


# --------------------------------------------------------------------------- #
# Plain text / Markdown
# --------------------------------------------------------------------------- #
_MD_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_MD_SETEXT = re.compile(r"^(=+|-+)\s*$")


def _parse_text(data: bytes, is_markdown: bool) -> ParsedDocument:
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1", errors="replace")
    if not text.strip():
        raise ValueError("Text file is empty.")

    text = clean_text(text)
    lines = text.splitlines()
    blocks: List[Block] = []
    buffer: List[str] = []
    page, chars = 1, 0

    def flush():
        nonlocal buffer
        if buffer:
            blocks.append(Block(text="\n".join(buffer), page=page))
            buffer = []

    for i, line in enumerate(lines):
        stripped = line.strip()

        heading = _MD_HEADING.match(stripped) if is_markdown else None
        setext = (
            is_markdown and stripped and i + 1 < len(lines) and _MD_SETEXT.match(lines[i + 1].strip())
        )
        if setext and _MD_SETEXT.match(stripped):
            continue                                   # the underline itself

        if heading:
            flush()
            blocks.append(Block(text=heading.group(2).strip(), page=page,
                                kind="heading", level=len(heading.group(1))))
        elif setext:
            flush()
            level = 1 if lines[i + 1].strip().startswith("=") else 2
            blocks.append(Block(text=stripped, page=page, kind="heading", level=level))
        elif not stripped:
            flush()
        else:
            buffer.append(stripped)

        chars += len(line)
        if chars > 3000:
            flush()
            page += 1
            chars = 0
    flush()

    if not blocks:
        raise ValueError("Text file has no usable content.")

    blocks = _assign_sections(blocks)
    # No heading anywhere -- leave the title empty so the caller can fall back to
    # the filename, which reads far better than the first line of prose.
    title = next((b.text for b in blocks if b.kind == "heading"), "")
    return ParsedDocument(blocks=blocks, title=title, page_count=page,
                          meta={"format": "md" if is_markdown else "txt"})


# --------------------------------------------------------------------------- #
# Entry point
# --------------------------------------------------------------------------- #
def parse_document(file_bytes: bytes, file_type: str, filename: str = "") -> ParsedDocument:
    """Parse raw bytes into a structured document. Raises ValueError on bad input."""
    file_type = file_type.lower().lstrip(".")
    if file_type not in SUPPORTED_TYPES:
        raise ValueError(f"Unsupported type '{file_type}'. Allowed: {', '.join(sorted(SUPPORTED_TYPES))}")
    if not file_bytes:
        raise ValueError("File is empty.")

    if file_type == "pdf":
        parsed = _parse_pdf(file_bytes)
    elif file_type == "docx":
        parsed = _parse_docx(file_bytes)
    else:
        parsed = _parse_text(file_bytes, is_markdown=file_type in ("md", "markdown"))

    if not parsed.title and not filename and parsed.blocks:
        parsed.title = parsed.blocks[0].text[:80]
    if filename and (not parsed.title or len(parsed.title) < 3):
        parsed.title = filename.rsplit(".", 1)[0].replace("_", " ").replace("-", " ").strip()
    parsed.meta["filename"] = filename
    return parsed
