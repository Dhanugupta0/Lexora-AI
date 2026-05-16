"""Document parser — supports PDF (PyMuPDF), DOCX, and plain text."""
import io
import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import List

logger = logging.getLogger(__name__)

SUPPORTED_TYPES = {"pdf", "docx", "txt"}


@dataclass
class ParsedDocument:
    """Output of the parser: extracted text split by page."""

    pages: List[str] = field(default_factory=list)  # one entry per page/section

    @property
    def page_count(self) -> int:
        return len(self.pages)

    @property
    def full_text(self) -> str:
        return "\n\n".join(self.pages)


class ParserError(Exception):
    pass


def parse_document(file_bytes: bytes, file_type: str) -> ParsedDocument:
    """Parse raw file bytes into a :class:`ParsedDocument`.

    Args:
        file_bytes: Raw content of the uploaded file.
        file_type: Lowercase extension without dot (pdf | docx | txt).

    Returns:
        ParsedDocument with one list entry per page/section.

    Raises:
        ParserError: If parsing fails or file type is unsupported.
    """
    file_type = file_type.lower().lstrip(".")
    if file_type not in SUPPORTED_TYPES:
        raise ParserError(f"Unsupported file type: '{file_type}'. Supported: {SUPPORTED_TYPES}")

    try:
        if file_type == "pdf":
            return _parse_pdf(file_bytes)
        elif file_type == "docx":
            return _parse_docx(file_bytes)
        else:
            return _parse_txt(file_bytes)
    except ParserError:
        raise
    except Exception as exc:
        logger.exception("Unexpected error while parsing %s document", file_type)
        raise ParserError(f"Failed to parse {file_type} document: {exc}") from exc


def _parse_pdf(data: bytes) -> ParsedDocument:
    """Extract text page-by-page using PyMuPDF."""
    import fitz  # PyMuPDF

    pages: List[str] = []
    with fitz.open(stream=data, filetype="pdf") as doc:
        for page in doc:
            text = page.get_text("text")
            if text.strip():
                pages.append(text)
    if not pages:
        raise ParserError("PDF contains no extractable text (may be a scanned image).")
    return ParsedDocument(pages=pages)


def _parse_docx(data: bytes) -> ParsedDocument:
    """Extract paragraph text from a DOCX file using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(data))
    sections: List[str] = []
    current: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            if current:
                sections.append("\n".join(current))
                current = []
        else:
            current.append(text)

    if current:
        sections.append("\n".join(current))

    if not sections:
        raise ParserError("DOCX contains no extractable text.")
    return ParsedDocument(pages=sections)


def _parse_txt(data: bytes) -> ParsedDocument:
    """Split plain text into ~50-line virtual pages."""
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1")

    lines = text.splitlines()
    page_size = 50
    pages: List[str] = []
    for i in range(0, max(len(lines), 1), page_size):
        chunk = "\n".join(lines[i : i + page_size]).strip()
        if chunk:
            pages.append(chunk)

    if not pages:
        raise ParserError("Text file is empty.")
    return ParsedDocument(pages=pages)
